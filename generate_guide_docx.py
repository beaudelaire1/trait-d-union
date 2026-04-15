"""Génère le Guide Diagnostic Client TUS en .docx"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── Couleurs TUS ──
TUS_BLACK = RGBColor(0x07, 0x08, 0x0A)
TUS_BLUE = RGBColor(0x0B, 0x2D, 0xFF)
TUS_GREEN = RGBColor(0x22, 0xC5, 0x5E)
TUS_RED = RGBColor(0xDC, 0x26, 0x26)
TUS_ORANGE = RGBColor(0xF9, 0x73, 0x16)
TUS_YELLOW = RGBColor(0xEA, 0xB3, 0x08)
TUS_GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Styles de base ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = TUS_BLACK

for i in range(1, 5):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = TUS_BLUE if i <= 2 else TUS_BLACK
    h.font.name = 'Calibri'
    h.font.bold = True
    if i == 1:
        h.font.size = Pt(22)
    elif i == 2:
        h.font.size = Pt(16)
    elif i == 3:
        h.font.size = Pt(13)
    else:
        h.font.size = Pt(11)

# ── Marges ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_table(headers, rows, col_widths=None):
    """Ajoute un tableau stylisé."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Blue background
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '0B2DFF',
            qn('w:val'): 'clear',
        })
        shading.append(bg)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            # Alternate row shading
            if r_idx % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                bg = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F4FF',
                    qn('w:val'): 'clear',
                })
                shading.append(bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return table


def add_quote(text):
    """Ajoute une citation stylisée."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'« {text} »')
    run.italic = True
    run.font.color.rgb = TUS_GRAY
    run.font.size = Pt(10)


def add_checklist(items):
    """Ajoute une liste de cases à cocher."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run('☐  ')
        run.font.size = Pt(10)
        run = p.add_run(item)
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'


def add_bullet(text, bold_prefix=None):
    """Ajoute un bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.size = Pt(10)
    else:
        p.runs[0].font.size = Pt(10) if p.runs else None
        if not p.runs:
            run = p.add_run(text)
            run.font.size = Pt(10)
        else:
            p.text = text


# ═══════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TRAIT D\'UNION STUDIO')
run.font.size = Pt(14)
run.font.color.rgb = TUS_BLUE
run.bold = True
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Guide de Diagnostic Client')
run.font.size = Pt(28)
run.font.color.rgb = TUS_BLACK
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Protocole opérationnel pour réaliser un diagnostic complet\nde l\'activité et/ou du site web d\'un client')
run.font.size = Pt(12)
run.font.color.rgb = TUS_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Avril 2026')
run.font.size = Pt(11)
run.font.color.rgb = TUS_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Durée estimée : 2 à 4 heures (complet) · 1 heure (ciblé)')
run.font.size = Pt(10)
run.font.color.rgb = TUS_GRAY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Table des matières', level=1)

toc_items = [
    'Phase 0 — Prise de contact & qualification',
    'Phase 1 — Entretien de découverte',
    'Phase 2 — Diagnostic Activité (les 29 simulateurs TUS)',
    'Phase 3 — Diagnostic Site Web (sans accès au code source)',
    'Phase 4 — Synthèse & scoring',
    'Phase 5 — Livrable & proposition',
    'Annexes — Outils externes & checklist terrain',
]
for i, item in enumerate(toc_items):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {item}')
    run.font.size = Pt(11)
    if i == 0:
        run.text = f'{item}'
    run.text = f'{i}. {item}'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PHASE 0
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Phase 0 — Prise de contact & qualification', level=1)

p = doc.add_paragraph('Avant tout diagnostic, qualifier le besoin du client :')

add_table(
    ['Question', 'Pourquoi'],
    [
        ['Quel est votre secteur d\'activité ?', 'Adapter les simulateurs pertinents'],
        ['Avez-vous un site web existant ? Si oui, l\'URL ?', 'Savoir si on fait un diagnostic site en plus'],
        ['Sur quelle plateforme est-il construit ?', 'Adapter les outils d\'analyse'],
        ['Quel est votre objectif principal ?', 'Cadrer le périmètre du diagnostic'],
        ['Depuis combien de temps l\'activité existe ?', 'Contextualiser les chiffres'],
        ['Êtes-vous seul ou avez-vous une équipe ?', 'Impact sur les simulateurs RH/délégation'],
    ],
    col_widths=[3.0, 3.5],
)

p = doc.add_paragraph()
run = p.add_run('Résultat : ')
run.bold = True
run = p.add_run('on sait si on fait un diagnostic activité, un diagnostic site, ou les deux.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PHASE 1
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Phase 1 — Entretien de découverte', level=1)

p = doc.add_paragraph(
    'Entretien structuré de 30 à 45 minutes avec le client. '
    'Ne pas envoyer de questionnaire froid — c\'est une conversation '
    'qui fait partie de l\'expérience TUS.'
)

doc.add_heading('Données à recueillir — Activité', level=2)

add_table(
    ['Donnée', 'Exemple', 'Utilisée dans'],
    [
        ['CA mensuel moyen', '12 000 €', 'Seuil de rentabilité, Atterrissage, Plafond'],
        ['Charges fixes mensuelles', '6 500 €', 'Seuil de rentabilité, Couverture CF'],
        ['Charges variables (% du CA)', '25%', 'Seuil de rentabilité, Mix produits'],
        ['Nombre de clients actifs', '18', 'Dépendance commerciale, Jumeaux clients'],
        ['Panier moyen', '650 €', 'CAC, Élasticité-prix, Prix psychologique'],
        ['Devis envoyés / mois', '8', 'Taux de conversion'],
        ['Devis signés / mois', '3', 'Taux de conversion'],
        ['Budget marketing mensuel', '400 €', 'CAC, ROI marketing'],
        ['Nouveaux clients / mois', '2', 'CAC, Rétention'],
        ['Délai moyen paiement client', '38 jours', 'DSO / trésorerie'],
        ['% CA récurrent vs ponctuel', '30/70', 'Stabilité du modèle'],
        ['Heures travaillées / semaine', '55h', 'Capacité facturable'],
        ['Heures facturées / semaine', '25h', 'Capacité facturable'],
        ['Nombre d\'outils SaaS', '12', 'Fragmentation'],
        ['Part du plus gros client', '40%', 'Dépendance commerciale'],
        ['Nombre de prestations différentes', '6', 'Mix produits, Corrélation'],
        ['Effectif', '3', 'Seuil de délégation'],
        ['Saisonnalité perçue', '"creux janv-fév"', 'Saisonnalité'],
    ],
    col_widths=[2.2, 1.3, 3.0],
)

doc.add_heading('Données à recueillir — Site web', level=2)

add_table(
    ['Donnée', 'À demander au client'],
    [
        ['URL du site', 'www.exemple.gf'],
        ['Objectif du site', 'Vitrine ? Prise de RDV ? Vente en ligne ?'],
        ['Qui l\'a construit et quand ?', 'Agence, freelance, lui-même'],
        ['Google Analytics / Search Console ?', 'Si oui, accès lecteur possible ?'],
        ['Nombre de visiteurs / mois', 'Approximation OK'],
        ['D\'où viennent ses clients ?', 'Bouche à oreille, Google, réseaux…'],
        ['Publicité en ligne déjà faite ?', 'Google Ads, Facebook Ads, etc.'],
        ['Concurrents directs ?', 'URL si possible pour benchmark'],
    ],
    col_widths=[2.5, 4.0],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PHASE 2
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Phase 2 — Diagnostic Activité', level=1)

doc.add_heading('Les 6 indicateurs vitaux', level=2)

p = doc.add_paragraph(
    'Avant de lancer les simulateurs, poser les 6 KPIs de santé '
    'qui forment le socle du diagnostic :'
)

add_table(
    ['#', 'Indicateur', 'Calcul rapide', 'Zone danger'],
    [
        ['1', 'Trésorerie prévisionnelle à 30j', 'Solde + encaissements − décaissements', 'Négatif'],
        ['2', 'DSO (délai paiement client)', '(Créances / CA) × 30', '> 45 jours'],
        ['3', 'Taux de marge brute', '(CA − coûts directs) / CA', 'En baisse sur 2 mois'],
        ['4', 'Taux de conversion', 'Devis signés / devis envoyés', '< 25%'],
        ['5', 'Répartition CA récurrent / ponctuel', '% de CA sous contrat', '< 20% récurrent'],
        ['6', 'Couverture des charges fixes', 'Charges fixes / marge brute', '> 85%'],
    ],
    col_widths=[0.3, 2.0, 2.5, 1.5],
)

doc.add_heading('Règle des signaux simultanés', level=2)

add_table(
    ['Situation', 'Niveau', 'Action'],
    [
        ['1 indicateur en zone danger', '⚠️ Avertissement', 'Surveiller, mentionner dans le rapport'],
        ['2 indicateurs dégradés', '🚨 Alerte', 'Recommandation d\'action prioritaire'],
        ['3+ indicateurs dégradés', '🚨🚨🚨 Urgence', 'Réunion stratégique dans la semaine'],
    ],
    col_widths=[2.2, 1.5, 2.8],
)

doc.add_heading('Sélection des simulateurs selon le profil', level=2)

# Profil A
doc.add_heading('Profil A — TPE solo / freelance', level=3)

add_table(
    ['Simulateur TUS', 'Ce qu\'on cherche'],
    [
        ['Seuil de Rentabilité', 'Quel CA minimum pour couvrir les charges ?'],
        ['Capacité Maximale Facturable', 'Combien peut-il facturer au max ?'],
        ['Plafond de Verre', 'À quel CA l\'activité va plafonner ?'],
        ['Friction Opérationnelle', 'Combien coûtent le chaos et les process cassés ?'],
        ['Indice de Fragmentation', 'Trop d\'outils SaaS ? Lesquels éliminer ?'],
        ['Élasticité-Prix', 'Peut-il augmenter ses tarifs ?'],
        ['Prix Psychologique', 'Où est le prix perçu comme juste ?'],
        ['Coût d\'Inaction', 'Combien coûte le fait de ne rien changer ?'],
        ['Vallée de la Mort', 'Phase critique de croissance ?'],
    ],
    col_widths=[2.5, 4.0],
)

# Profil B
doc.add_heading('Profil B — PME 3-15 personnes', level=3)

add_table(
    ['Simulateur TUS', 'Ce qu\'on cherche'],
    [
        ['Seuil de Rentabilité', 'Point mort mensuel'],
        ['Atterrissage Trimestriel / Annuel', 'Va-t-on atteindre l\'objectif ?'],
        ['Atterrissage Mensuel Trésorerie', 'Tension de trésorerie imminente ?'],
        ['Radar de Dépendance Commerciale', 'Un client = trop du CA ?'],
        ['Seuil de Délégation', 'Embaucher ou externaliser ?'],
        ['Optimiseur de Mix Produits', 'Quelle prestation prioriser ?'],
        ['Jumeaux Clients', 'Quels segments sont les plus rentables ?'],
        ['Corrélation Produits/Services', 'Quels produits s\'achètent ensemble ?'],
        ['Impact Saisonnalité', 'Comment lisser les creux ?'],
        ['Coût de la Non-Qualité', 'Combien coûtent retours, SAV, erreurs ?'],
        ['Matrice Effort / Impact', 'Quels projets lancer en priorité ?'],
    ],
    col_widths=[2.5, 4.0],
)

# Profil C
doc.add_heading('Profil C — Réflexion stratégique', level=3)

add_table(
    ['Simulateur TUS', 'Ce qu\'on cherche'],
    [
        ['Scénario Pivot', 'Et si on changeait de modèle ?'],
        ['Taille de Marché Accessible', 'Le marché local est-il assez grand ?'],
        ['Valeur de Sortie', 'Combien vaut l\'entreprise ?'],
        ['Pricing par Paliers', 'Comment structurer l\'offre en gammes ?'],
        ['Vulnérabilité Fournisseur', 'Dépendance critique ?'],
        ['ROI Campagne Marketing', 'Les pubs ont-elles rapporté ?'],
        ['Vrai Coût d\'une Promotion', 'Ce -20%, ça coûte combien vraiment ?'],
    ],
    col_widths=[2.5, 4.0],
)

# Liste complète
doc.add_heading('Liste complète des 29 simulateurs TUS', level=3)

outils_historiques = [
    'Seuil de Rentabilité (Point Mort)',
    'Coût d\'Acquisition (CAC)',
    'Friction Opérationnelle',
    'Indice de Fragmentation',
    'Flux A.C.S.E',
    'Plafond de Verre',
    'Élasticité-Prix',
    'Vallée de la Mort',
    'Effet Rétention',
]
outils_nouveaux = [
    'Optimiseur de Mix Produits/Services',
    'Atterrissage Trimestriel / Annuel',
    'Atterrissage Mensuel Trésorerie',
    'Simulateur de Jumeaux Clients',
    'Corrélation Produits/Services',
    'Seuil de Délégation',
    'Prix Psychologique',
    'Radar de Dépendance Commerciale',
    'Capacité Maximale Facturable',
    'Impact Saisonnalité',
    'Vrai Coût d\'une Promotion',
    'Valeur de Sortie',
    'Matrice Effort / Impact',
    'Coût d\'Inaction',
    'Scénario Pivot',
    'ROI Campagne Marketing',
    'Pricing par Paliers',
    'Taille de Marché Accessible',
    'Vulnérabilité Fournisseur',
    'Coût de la Non-Qualité',
]

p = doc.add_paragraph()
run = p.add_run('Outils historiques (9)')
run.bold = True
for i, o in enumerate(outils_historiques, 1):
    p = doc.add_paragraph(f'{i}. {o}')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
run = p.add_run('Nouveaux outils (20)')
run.bold = True
for i, o in enumerate(outils_nouveaux, 10):
    p = doc.add_paragraph(f'{i}. {o}')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PHASE 3
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Phase 3 — Diagnostic Site Web (sans accès au code source)', level=1)

p = doc.add_paragraph()
run = p.add_run('Prérequis : ')
run.bold = True
run = p.add_run('uniquement l\'URL du site. Aucun accès au code source nécessaire.')

# 3.1 Performance
doc.add_heading('3.1 — Performance & Vitesse', level=2)

add_table(
    ['Outil', 'URL', 'Gratuit', 'Ce qu\'on mesure'],
    [
        ['Google PageSpeed Insights', 'pagespeed.web.dev', '✅', 'Core Web Vitals : LCP, INP, CLS'],
        ['GTmetrix', 'gtmetrix.com', '✅', 'Temps de chargement, waterfall'],
        ['WebPageTest', 'webpagetest.org', '✅', 'TTFB, First Byte, Speed Index'],
        ['Pingdom Tools', 'tools.pingdom.com', '✅', 'Temps de réponse, poids page'],
    ],
    col_widths=[1.8, 1.6, 0.6, 2.5],
)

p = doc.add_paragraph()
run = p.add_run('Points à noter dans le rapport :')
run.bold = True

add_checklist([
    'Score PageSpeed Mobile (objectif : > 70)',
    'Score PageSpeed Desktop (objectif : > 85)',
    'LCP — Largest Contentful Paint (objectif : < 2.5s)',
    'INP — Interaction to Next Paint (objectif : < 200ms)',
    'CLS — Cumulative Layout Shift (objectif : < 0.1)',
    'TTFB — Time To First Byte (objectif : < 800ms)',
    'Poids total de la page (objectif : < 3 Mo)',
    'Nombre de requêtes HTTP (objectif : < 50)',
    'Images non optimisées (format, compression, dimensions)',
    'Cache navigateur configuré ?',
    'Compression GZIP/Brotli activée ?',
])

# 3.2 SEO
doc.add_heading('3.2 — SEO (Référencement naturel)', level=2)

add_table(
    ['Outil', 'URL', 'Gratuit', 'Ce qu\'on mesure'],
    [
        ['Google Search Console', 'search.google.com/search-console', '✅', 'Indexation, erreurs, requêtes'],
        ['Screaming Frog', 'screamingfrog.co.uk', '✅ < 500 URLs', 'Crawl complet : titres, méta, H1, 404'],
        ['Ahrefs Webmaster Tools', 'ahrefs.com/webmaster-tools', '✅', 'Backlinks, santé technique'],
        ['Ubersuggest', 'neilpatel.com/ubersuggest', 'Freemium', 'Mots-clés, volume, audit SEO'],
        ['Rich Results Test', 'search.google.com/test/rich-results', '✅', 'Données structurées (schema.org)'],
    ],
    col_widths=[1.6, 2.3, 0.9, 1.7],
)

p = doc.add_paragraph()
run = p.add_run('Points à noter dans le rapport :')
run.bold = True

add_checklist([
    'Le site est-il indexé par Google ? (site:domaine.com)',
    'Nombre de pages indexées vs pages réelles',
    'Chaque page a un <title> unique et pertinent ? (< 60 car.)',
    'Chaque page a une <meta description> unique ? (< 160 car.)',
    'Structure des titres correcte ? (un seul H1, hiérarchie H2 > H3)',
    'Images avec attribut alt renseigné ?',
    'URLs propres et lisibles ? (pas de ?p=123)',
    'Sitemap XML présent et soumis ?',
    'Fichier robots.txt présent et correct ?',
    'Liens internes cohérents ? Liens cassés (404) ?',
    'Redirections 301 en place si refonte précédente ?',
    'Données structurées (LocalBusiness, FAQ, Breadcrumb…) ?',
    'Mobile-first indexing OK ?',
])

# 3.3 Sécurité
doc.add_heading('3.3 — Sécurité (sans accès au code)', level=2)

add_table(
    ['Outil', 'URL', 'Gratuit', 'Ce qu\'on vérifie'],
    [
        ['SSL Labs', 'ssllabs.com/ssltest', '✅', 'Certificat HTTPS, config TLS, note A-F'],
        ['Security Headers', 'securityheaders.com', '✅', 'CSP, HSTS, X-Frame-Options'],
        ['Mozilla Observatory', 'observatory.mozilla.org', '✅', 'Score sécurité global'],
        ['Sucuri SiteCheck', 'sitecheck.sucuri.net', '✅', 'Malware, blacklist, CMS obsolète'],
        ['WPScan (si WordPress)', 'wpscan.com', 'Freemium', 'Vulnérabilités plugins/thème'],
    ],
    col_widths=[1.6, 1.8, 0.8, 2.3],
)

p = doc.add_paragraph()
run = p.add_run('Points à noter dans le rapport :')
run.bold = True

add_checklist([
    'HTTPS actif ? Certificat valide et non expiré ?',
    'Note SSL Labs (objectif : A ou A+)',
    'Redirection HTTP → HTTPS en place ?',
    'En-tête HSTS activé ?',
    'En-tête Content-Security-Policy présent ?',
    'En-tête X-Frame-Options présent ?',
    'Version du CMS exposée publiquement ?',
    'Plugins/extensions avec vulnérabilités connues ?',
    'Page de login par défaut accessible ? (/wp-admin, /wp-login.php)',
    'Fichiers sensibles accessibles ? (/readme.html, /xmlrpc.php, /.env)',
    'Le site est-il sur des listes noires ?',
])

# 3.4 Accessibilité
doc.add_heading('3.4 — Accessibilité (WCAG)', level=2)

add_table(
    ['Outil', 'URL', 'Gratuit', 'Ce qu\'on mesure'],
    [
        ['WAVE', 'wave.webaim.org', '✅', 'Erreurs WCAG, contrastes, ARIA'],
        ['axe DevTools', 'Extension navigateur', '✅', 'Audit WCAG automatisé'],
        ['Lighthouse', 'Chrome DevTools (F12)', '✅', 'Score accessibilité global'],
        ['Contrast Checker', 'webaim.org/resources/contrastchecker', '✅', 'Ratio de contraste texte/fond'],
    ],
    col_widths=[1.5, 2.3, 0.7, 2.0],
)

p = doc.add_paragraph()
run = p.add_run('Points à noter dans le rapport :')
run.bold = True

add_checklist([
    'Score Lighthouse Accessibilité (objectif : > 85)',
    'Nombre d\'erreurs WAVE (objectif : 0 erreur critique)',
    'Navigation au clavier fonctionnelle ? (Tab, Enter, Escape)',
    'Focus visible sur les éléments interactifs ?',
    'Contrastes texte/fond suffisants ? (ratio 4.5:1 minimum)',
    'Textes alternatifs sur toutes les images ?',
    'Formulaires avec labels associés ?',
    'Langue de la page déclarée (<html lang="fr">) ?',
    'Hiérarchie des titres logique ?',
])

# 3.5 Mobile
doc.add_heading('3.5 — Mobile & Responsive', level=2)

add_table(
    ['Outil', 'URL', 'Gratuit', 'Ce qu\'on mesure'],
    [
        ['Google Mobile-Friendly Test', 'search.google.com/test/mobile-friendly', '✅', 'Compatibilité mobile'],
        ['Responsively App', 'responsively.app', '✅', 'Prévisualisation multi-écrans'],
        ['Chrome DevTools', 'F12 > Toggle device', '✅', 'Simulation de résolutions'],
    ],
    col_widths=[1.8, 2.5, 0.7, 1.5],
)

p = doc.add_paragraph()
run = p.add_run('Points à noter dans le rapport :')
run.bold = True

add_checklist([
    'Le site est-il responsive ?',
    'Texte lisible sans zoom sur mobile ?',
    'Boutons/liens assez espacés pour le tactile ? (min 44×44 px)',
    'Pas de défilement horizontal ?',
    'Menu de navigation utilisable sur mobile ?',
    'Formulaires faciles à remplir sur mobile ?',
    'Images qui s\'adaptent (pas de débordement) ?',
    'Viewport meta tag présent ?',
])

# 3.6 UX
doc.add_heading('3.6 — UX / Design / Contenu (inspection manuelle)', level=2)

p = doc.add_paragraph()
run = p.add_run('Pas d\'outil automatisé — c\'est l\'œil du professionnel qui parle.')
run.italic = True

doc.add_heading('Première impression (test des 5 secondes)', level=3)
add_checklist([
    'En 5 secondes, comprend-on ce que fait l\'entreprise ?',
    'La proposition de valeur est-elle visible immédiatement ?',
    'Le CTA principal est-il évident ?',
])

doc.add_heading('Navigation & architecture', level=3)
add_checklist([
    'Le menu est-il clair et logique ? (< 7 items)',
    'Accès à n\'importe quelle page en 3 clics maximum ?',
    'Fil d\'Ariane présent ?',
    'Page 404 personnalisée ?',
    'Recherche interne disponible (si > 20 pages) ?',
])

doc.add_heading('Contenu', level=3)
add_checklist([
    'Textes clairs, sans jargon, orientés bénéfice client ?',
    'Orthographe et grammaire impeccables ?',
    'Témoignages / avis clients présents ?',
    'Preuves sociales (logos clients, certifications, chiffres clés) ?',
    'Visuels de qualité professionnelle ?',
    'Cohérence graphique (couleurs, typos, espacements) ?',
])

doc.add_heading('Conversion', level=3)
add_checklist([
    'Combien de clics pour contacter / demander un devis ? (objectif : ≤ 2)',
    'Numéro de téléphone cliquable sur mobile ?',
    'Formulaire de contact fonctionnel et simple ? (< 5 champs)',
    'CTA visibles et contrastés sur chaque page ?',
    'Réassurance présente ? (garanties, délais, processus)',
])

doc.add_heading('Conformité légale', level=3)
add_checklist([
    'Mentions légales présentes et complètes ?',
    'Politique de confidentialité / RGPD ?',
    'Bandeau cookies conforme ?',
    'CGV si e-commerce ?',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PHASE 4
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Phase 4 — Synthèse & scoring', level=1)

doc.add_heading('Grille de scoring', level=2)

add_table(
    ['Domaine', 'Poids', 'Note /10', 'Score pondéré'],
    [
        ['Santé financière (6 KPIs vitaux)', '25%', '___', '___'],
        ['Stratégie & positionnement (simulateurs)', '15%', '___', '___'],
        ['Performance site (vitesse, Core Web Vitals)', '15%', '___', '___'],
        ['SEO (référencement, indexation)', '15%', '___', '___'],
        ['Sécurité (HTTPS, headers, CMS)', '10%', '___', '___'],
        ['Mobile & responsive', '5%', '___', '___'],
        ['Accessibilité', '5%', '___', '___'],
        ['UX / Design / Contenu', '5%', '___', '___'],
        ['Conformité légale', '5%', '___', '___'],
        ['SCORE GLOBAL', '100%', '', '/10'],
    ],
    col_widths=[2.8, 0.7, 0.8, 1.2],
)

doc.add_heading('Barème', level=2)

add_table(
    ['Note', 'Signification'],
    [
        ['9-10', 'Excellent — rien à signaler, maintenir'],
        ['7-8', 'Bon — améliorations mineures possibles'],
        ['5-6', 'Passable — plusieurs points à corriger'],
        ['3-4', 'Faible — problèmes significatifs'],
        ['1-2', 'Critique — intervention urgente nécessaire'],
    ],
    col_widths=[1.0, 5.5],
)

doc.add_heading('Classification des recommandations', level=2)

add_table(
    ['Priorité', 'Critère', 'Exemple'],
    [
        ['🔴 Critique', 'Impact direct sur CA ou sécurité', 'Site non HTTPS, formulaire cassé, tréso négative'],
        ['🟠 Important', 'Frein significatif à la croissance', 'PageSpeed < 40, pas de SEO, dépendance 1 client > 50%'],
        ['🟡 Recommandé', 'Amélioration notable', 'Optimiser images, ajouter témoignages, données structurées'],
        ['🟢 Bonus', 'Finition, différenciation', 'Micro-animations, A/B testing, schema avancé'],
    ],
    col_widths=[1.0, 2.0, 3.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PHASE 5
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Phase 5 — Livrable & proposition', level=1)

doc.add_heading('Structure du rapport diagnostic', level=2)

structure = [
    ('1. Page de garde', 'Logo TUS + logo client, titre, date, nom du client'),
    ('2. Résumé exécutif (1 page)', 'Score global /10, 3 forces, 3 faiblesses, recommandation prioritaire n°1'),
    ('3. Diagnostic Activité', 'Tableau des 6 KPIs avec voyants, résultats simulateurs, signaux d\'alerte'),
    ('4. Diagnostic Site Web', 'Résultats par domaine, captures annotées, scores des outils'),
    ('5. Plan d\'action', 'Actions 🔴 Critiques → 🟠 Importantes → 🟡 Recommandées → 🟢 Bonus'),
    ('6. Proposition TUS', 'Services adaptés au diagnostic, devis personnalisé'),
]

for title, detail in structure:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    p = doc.add_paragraph(detail)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

doc.add_heading('Comment présenter le rapport', level=2)

regles = [
    ('Jamais par email seul', 'toujours en visio ou en présentiel'),
    ('Commencer par les forces', 'le client doit entendre ce qui marche avant ce qui ne marche pas'),
    ('Chiffrer l\'impact', 'utiliser le simulateur Coût d\'Inaction pour montrer combien l\'inaction coûte'),
    ('Proposer 3 niveaux d\'intervention', 'essentiel / recommandé / premium'),
]

for i, (titre, detail) in enumerate(regles, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {titre} — ')
    run.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run(detail)
    run.font.size = Pt(10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ANNEXES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Annexes', level=1)

# Annexe A
doc.add_heading('A. Checklist terrain express', level=2)

p = doc.add_paragraph()
run = p.add_run('CLIENT : ________________________  DATE : ___________  URL : ________________________')
run.font.name = 'Consolas'
run.font.size = Pt(9)

sections_checklist = {
    'ACTIVITÉ': [
        'CA mensuel : ________    Charges fixes : ________    Marge brute : ________%',
        'Tréso prévisionnelle 30j : ________    DSO : ________ jours',
        'Taux conversion : ________%    CA récurrent : ________%    Couverture CF : ________%',
        'Signaux simultanés :  ☐ 0   ☐ 1 ⚠️   ☐ 2 🚨   ☐ 3+ 🚨🚨🚨',
    ],
    'PERFORMANCE SITE': [
        'PageSpeed Mobile : ___/100    PageSpeed Desktop : ___/100',
        'LCP : ___s    INP : ___ms    CLS : ___    TTFB : ___ms',
        'Poids page : ___ Mo    Requêtes : ___',
    ],
    'SEO': [
        'Indexé Google : ☐ Oui ☐ Non    Pages indexées : ___',
        'Titles uniques : ☐ Oui ☐ Non    Méta desc : ☐ Oui ☐ Non',
        'H1 unique/page : ☐ Oui ☐ Non    Sitemap : ☐ Oui ☐ Non',
        'Alt images : ☐ Oui ☐ Non    Liens cassés : ___',
    ],
    'SÉCURITÉ': [
        'HTTPS : ☐ Oui ☐ Non    SSL Labs : ___    Security Headers : ___',
        'CMS à jour : ☐ Oui ☐ Non ☐ N/A    Infos sensibles exposées : ☐ Oui ☐ Non',
    ],
    'ACCESSIBILITÉ': [
        'Lighthouse : ___/100    Erreurs WAVE : ___    Contrastes OK : ☐ Oui ☐ Non',
        'Navigation clavier : ☐ Oui ☐ Non    lang="fr" : ☐ Oui ☐ Non',
    ],
    'MOBILE': [
        'Mobile-friendly : ☐ Oui ☐ Non    Texte lisible : ☐ Oui ☐ Non',
        'Boutons touch OK : ☐ Oui ☐ Non    Pas de scroll H : ☐ Oui ☐ Non',
    ],
    'UX / CONTENU': [
        'Proposition valeur claire : ☐ Oui ☐ Non    CTA visible : ☐ Oui ☐ Non',
        'Contact en ≤ 2 clics : ☐ Oui ☐ Non    Témoignages : ☐ Oui ☐ Non',
        'Mentions légales : ☐ Oui ☐ Non    RGPD/cookies : ☐ Oui ☐ Non',
    ],
}

for section_title, items in sections_checklist.items():
    p = doc.add_paragraph()
    run = p.add_run(section_title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = TUS_BLUE
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run('SCORE GLOBAL : ___/10')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = TUS_BLUE

# Annexe B
doc.add_page_break()
doc.add_heading('B. Outils externes — récapitulatif', level=2)

add_table(
    ['Outil', 'Catégorie', 'URL', 'Coût'],
    [
        ['Google PageSpeed', 'Performance', 'pagespeed.web.dev', 'Gratuit'],
        ['GTmetrix', 'Performance', 'gtmetrix.com', 'Gratuit'],
        ['WebPageTest', 'Performance', 'webpagetest.org', 'Gratuit'],
        ['Pingdom Tools', 'Performance', 'tools.pingdom.com', 'Gratuit'],
        ['Google Search Console', 'SEO', 'search.google.com/search-console', 'Gratuit'],
        ['Screaming Frog', 'SEO', 'screamingfrog.co.uk', 'Gratuit < 500'],
        ['Ahrefs Webmaster', 'SEO', 'ahrefs.com/webmaster-tools', 'Gratuit'],
        ['Ubersuggest', 'SEO', 'neilpatel.com/ubersuggest', 'Freemium'],
        ['Rich Results Test', 'SEO', 'search.google.com/test/rich-results', 'Gratuit'],
        ['SSL Labs', 'Sécurité', 'ssllabs.com/ssltest', 'Gratuit'],
        ['Security Headers', 'Sécurité', 'securityheaders.com', 'Gratuit'],
        ['Mozilla Observatory', 'Sécurité', 'observatory.mozilla.org', 'Gratuit'],
        ['Sucuri SiteCheck', 'Sécurité', 'sitecheck.sucuri.net', 'Gratuit'],
        ['WPScan', 'Sécurité (WP)', 'wpscan.com', 'Freemium'],
        ['WAVE', 'Accessibilité', 'wave.webaim.org', 'Gratuit'],
        ['axe DevTools', 'Accessibilité', 'Extension navigateur', 'Gratuit'],
        ['Lighthouse', 'Perf+Access+SEO', 'Chrome DevTools (F12)', 'Gratuit'],
        ['Contrast Checker', 'Accessibilité', 'webaim.org/resources/contrastchecker', 'Gratuit'],
        ['Mobile-Friendly Test', 'Mobile', 'search.google.com/test/mobile-friendly', 'Gratuit'],
        ['Responsively', 'Mobile', 'responsively.app', 'Gratuit'],
    ],
    col_widths=[1.3, 1.2, 2.5, 0.8],
)

# Annexe C
doc.add_heading('C. Correspondance problème → simulateur TUS', level=2)

correspondances = [
    ['Ne sait pas s\'il gagne de l\'argent', 'Seuil de Rentabilité'],
    ['Dépense en pub sans savoir si ça marche', 'CAC + ROI Campagne Marketing'],
    ['Perd du temps dans des process manuels', 'Friction Opérationnelle'],
    ['15 outils SaaS qui ne communiquent pas', 'Indice de Fragmentation'],
    ['Ne sait pas quelle prestation prioriser', 'Mix Produits + Effort/Impact'],
    ['Hésite à embaucher', 'Seuil de Délégation'],
    ['Ne sait pas fixer ses prix', 'Élasticité + Prix Psycho + Pricing Paliers'],
    ['1 client = 50%+ du CA', 'Radar Dépendance Commerciale'],
    ['Travaille 60h, facture 30%', 'Capacité Maximale Facturable'],
    ['Creux d\'activité saisonniers', 'Impact Saisonnalité'],
    ['Veut faire une promo mais hésite', 'Vrai Coût d\'une Promotion'],
    ['Pense à vendre / transmettre', 'Valeur de Sortie'],
    ['Plusieurs projets, ne sait pas lequel', 'Matrice Effort / Impact'],
    ['Repousse des décisions importantes', 'Coût d\'Inaction'],
    ['Envisage un changement de modèle', 'Scénario Pivot'],
    ['Veut structurer l\'offre en gammes', 'Pricing par Paliers'],
    ['Se demande si le marché est assez grand', 'Taille de Marché Accessible'],
    ['Dépend d\'un fournisseur unique', 'Vulnérabilité Fournisseur'],
    ['Beaucoup de SAV / retours / erreurs', 'Coût de la Non-Qualité'],
    ['Ne sait pas quels clients fidéliser', 'Jumeaux Clients + Effet Rétention'],
    ['Quels produits s\'achètent ensemble ?', 'Corrélation Produits/Services'],
    ['Croissance stagne sans raison', 'Plafond de Verre + Vallée de la Mort'],
    ['Process de vente lent/bloqué', 'Flux A.C.S.E'],
    ['Ne sait pas s\'il atteindra l\'objectif', 'Atterrissage Trimestriel / Annuel'],
    ['Peur de ne pas passer le mois', 'Atterrissage Mensuel Trésorerie'],
]

add_table(
    ['Problème identifié', 'Simulateur(s) TUS'],
    correspondances,
    col_widths=[3.2, 3.3],
)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Trait d\'Union Studio — Avril 2026')
run.font.size = Pt(9)
run.font.color.rgb = TUS_GRAY
run.italic = True

# ── Save ──
output_path = r'c:\Users\vilme\OneDrive\Bureau\ressources_websites\trait_d_union_studio\src\GUIDE_DIAGNOSTIC_CLIENT.docx'
doc.save(output_path)
print(f'✅ Guide généré : {output_path}')
